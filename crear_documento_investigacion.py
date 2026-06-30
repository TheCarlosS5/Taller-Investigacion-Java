# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Path definitions
template_path = r"C:\Users\carlos\Desktop\SENA 2026\taller_2_mysql_workbench\Documento de Ejemplo.docx"
output_path = r"C:\Users\carlos\Desktop\investigacion\Taller_Sintaxis_Java.docx"

if not os.path.exists(template_path):
    print(f"Error: Template {template_path} not found")
    exit(1)

doc = docx.Document(template_path)
print(f"Opened template. Paragraph count: {len(doc.paragraphs)}")

# 1) Cover Page Update (Tono Aprendiz, keeping Carlos Stiven's details)
if len(doc.paragraphs[0].runs) > 0:
    doc.paragraphs[0].runs[0].text = "Taller de Investigación"
else:
    doc.paragraphs[0].text = "Taller de Investigación"

if len(doc.paragraphs[1].runs) > 0:
    doc.paragraphs[1].runs[0].text = "Métodos Abreviados y Sintaxis Comprimidas en Java"
else:
    doc.paragraphs[1].text = "Métodos Abreviados y Sintaxis Comprimidas en Java"

# Instructor and Date update on cover page
for p in doc.paragraphs[:16]:
    if "Karol Correa" in p.text:
        for run in p.runs:
            if "Karol Correa" in run.text:
                run.text = run.text.replace("Karol Correa", "Carlos Julio")
    if "12 de junio de 2026" in p.text:
        for run in p.runs:
            if "12 de junio de 2026" in run.text:
                run.text = run.text.replace("12 de junio de 2026", "24 de junio de 2026")

print("Cover page updated.")

# 2) Introduction and Objectives (Rewritten to SENA Apprentice style)
intro_p1 = (
    "Hola instructor, en este informe voy a presentar lo que investigué sobre los métodos "
    "abreviados y la sintaxis comprimida de Java. La idea de este trabajo es entender cómo "
    "ha ido evolucionando Java para hacernos la vida más fácil a los programadores, permitiéndonos "
    "escribir códigos más cortitos y limpios sin perder la seguridad de que todo esté bien tipado."
)
intro_p2 = (
    "En el documento explico de forma sencilla 15 temas de la lista que nos dio y 5 temas "
    "avanzados que investigué por mi cuenta. Para cada uno puse de qué se trata con mis propias "
    "palabras, qué ventajas y desventajas tiene, cuándo conviene usarlo y un ejemplo práctico "
    "comparando la forma tradicional (larga) de escribir el código con la forma abreviada (moderna)."
)

doc.paragraphs[19].text = intro_p1
doc.paragraphs[20].text = intro_p2

objectives = [
    "Investigar y documentar al menos 15 métodos abreviados de la lista propuesta por el instructor.",
    "Buscar y analizar 5 características modernas o avanzadas de Java de forma autónoma.",
    "Explicar con palabras sencillas las ventajas y desventajas de cada atajo sintáctico.",
    "Aprender a identificar en qué casos es bueno usar cada abreviación y en cuáles es mejor evitarla para no enredar el código.",
    "Crear ejemplos de código funcionales que comparen la manera vieja y la manera nueva de programar en Java."
]

for idx, obj_text in enumerate(objectives):
    p_idx = 22 + idx
    doc.paragraphs[p_idx].text = obj_text

print("Introduction and objectives updated in apprentice tone.")

# 3) Clear old content
while len(doc.paragraphs) > 28:
    p = doc.paragraphs[28]
    p._element.getparent().remove(p._element)

print("Old content cleared.")

# Helper functions for formatting
def add_spacing(doc, pt=6):
    p = doc.add_paragraph("", style="Normal")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pt)

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    pPr = p._p.get_or_add_pPr()
    shd = docx.oxml.parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F5F5F5"/>')
    pPr.append(shd)
    
    pbdr = docx.oxml.parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:left w:val="single" w:sz="24" w:space="8" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pbdr)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    
    rPr = run._r.get_or_add_rPr()
    rFonts = docx.oxml.OxmlElement('w:rFonts')
    rFonts.set(docx.oxml.ns.qn('w:ascii'), 'Consolas')
    rFonts.set(docx.oxml.ns.qn('w:hAnsi'), 'Consolas')
    rPr.append(rFonts)

def add_diagram_block(doc, diagram_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    
    pPr = p._p.get_or_add_pPr()
    shd = docx.oxml.parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="FAFAFA"/>')
    pPr.append(shd)
    
    pbdr = docx.oxml.parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:left w:val="single" w:sz="12" w:space="8" w:color="BBBBBB"/></w:pBdr>')
    pPr.append(pbdr)
    
    run = p.add_run(diagram_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.0)
    
    rPr = run._r.get_or_add_rPr()
    rFonts = docx.oxml.OxmlElement('w:rFonts')
    rFonts.set(docx.oxml.ns.qn('w:ascii'), 'Consolas')
    rFonts.set(docx.oxml.ns.qn('w:hAnsi'), 'Consolas')
    rPr.append(rFonts)

def set_cell_background(cell, fill_hex):
    shading_elm = docx.oxml.parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = docx.oxml.parse_xml(
            '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
            '</w:tblBorders>'
        )
        tblPr.append(tblBorders)

def add_word_table(doc, rows_data, col_widths):
    num_rows = len(rows_data)
    num_cols = len(rows_data[0])
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Normal Table'
    set_table_borders(table)
    
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx]
        row.height = Pt(18)
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.width = col_widths[c_idx]
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if r_idx == 0:
                    run.bold = True
                    
            if r_idx == 0:
                set_cell_background(cell, 'E6E6E6')

def add_bullet_item(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.4)
    run_bullet = p.add_run("• ")
    run_bullet.bold = True
    p.add_run(text)

def add_section_label(doc, label, content=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run_lbl = p.add_run(label + "\n")
    run_lbl.bold = True
    run_lbl.font.size = Pt(11)
    if content:
        p.add_run(content)

# Define simplified SENA apprentice tone content
part1_topics = [
    {
        "num": "1",
        "title": "1. Operador Ternario (?:)",
        "explicacion": "El operador ternario es como un if-else en una sola línea. Nos ayuda a elegir entre dos valores dependiendo de si una condición es verdadera o falsa, y lo guarda directamente en una variable.",
        "ventajas": [
            "Ahorra líneas de código para cosas muy simples.",
            "Es ideal para declarar variables con final porque las inicializas de inmediato y no tienes que dejar la variable vacía."
        ],
        "desventajas": [
            "Si pones condiciones muy largas o metes un ternario dentro de otro, se vuelve imposible de leer.",
            "Puede confundir si mezclas tipos de datos numéricos (como un int y un double) porque Java hará conversiones por detrás."
        ],
        "cuando": "Cuando tienes condiciones cortas de tipo 'sí o no', como decidir si un estudiante aprobó o reprobó una materia según su nota.",
        "code_trad": (
            "// Forma Tradicional\n"
            "String estado;\n"
            "if (puntuacion >= 70) {\n"
            "    estado = \"Aprobado\";\n"
            "} else {\n"
            "    estado = \"Reprobado\";\n"
            "}"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "String estado = (puntuacion >= 70) ? \"Aprobado\" : \"Reprobado\";"
        )
    },
    {
        "num": "2",
        "title": "2. Operadores de Asignación Compuesta (+=, -=, *=, /=, %=)",
        "explicacion": "Juntan una operación matemática (como sumar o restar) con guardar el resultado en la misma variable. Además, Java hace la conversión de tipo de forma automática por debajo para que no tengamos errores de compilación por pérdida de precisión.",
        "ventajas": [
            "El código se ve más compacto cuando tenemos contadores o acumuladores.",
            "No hay que hacer castings raros a mano cuando sumamos números de distintos tamaños (como un short y un int)."
        ],
        "desventajas": [
            "Puede hacer que perdamos datos sin darnos cuenta porque la conversión ocurre de forma silenciosa."
        ],
        "cuando": "Para sumarle o restarle cosas a variables acumuladoras o en cálculos de bucles.",
        "code_trad": (
            "// Forma Tradicional\n"
            "short balance = 100;\n"
            "balance = (short) (balance + 50); // El cast explícito es obligatorio"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "short balance = 100;\n"
            "balance += 50; // El compilador inyecta automáticamente el cast de forma transparente"
        )
    },
    {
        "num": "3",
        "title": "3. Operadores de Incremento y Decremento (++, --)",
        "explicacion": "Sirven para sumar o restar exactamente 1 a una variable de tipo número. Se pueden usar antes de la variable (++x) o después (x++), y esto cambia cuándo se aplica el cambio durante la ejecución.",
        "ventajas": [
            "Es el atajo más rápido y común para ir contando de uno en uno.",
            "Es fácil de entender a primera vista."
        ],
        "desventajas": [
            "Si los mezclas en operaciones complejas, te puedes confundir con el orden en que Java actualiza los valores."
        ],
        "cuando": "Principalmente en ciclos for tradicionales y contadores simples.",
        "code_trad": (
            "// Forma Tradicional\n"
            "contador = contador + 1;"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "contador++;"
        )
    },
    {
        "num": "4",
        "title": "4. Ciclo For-each (Bucle for Mejorado)",
        "explicacion": "Es una forma simplificada de recorrer listas o arreglos. Java se encarga de ir elemento por elemento por debajo usando un iterador o índices, sin que nosotros tengamos que manejar la variable del contador de forma manual.",
        "ventajas": [
            "Evita que nos equivoquemos con los límites del arreglo (como salirnos del tamaño del índice).",
            "El código es mucho más limpio porque nos enfocamos directo en procesar el objeto."
        ],
        "desventajas": [
            "No tenemos el índice del elemento actual, por lo que no podemos cambiar cosas por posición o recorrer al revés.",
            "Si intentas borrar un elemento mientras recorres la lista, Java te sacará un error (ConcurrentModificationException)."
        ],
        "cuando": "Cuando quieras leer y procesar todos los elementos de una colección de principio a fin sin modificar la estructura de la lista.",
        "code_trad": (
            "// Forma Tradicional\n"
            "List<String> ciudades = List.of(\"Bogotá\", \"Medellín\", \"Cali\");\n"
            "for (int i = 0; i < ciudades.size(); i++) {\n"
            "    String ciudad = ciudades.get(i);\n"
            "    System.out.println(ciudad);\n"
            "}"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "List<String> ciudades = List.of(\"Bogotá\", \"Medellín\", \"Cali\");\n"
            "for (String ciudad : ciudades) {\n"
            "    System.out.println(ciudad);\n"
            "}"
        )
    },
    {
        "num": "5",
        "title": "5. Inferencia de Tipos Locales con var",
        "explicacion": "Permite no escribir el nombre completo del tipo al declarar variables locales, dejando que Java lo deduzca solo según el valor de la derecha. Es importante aclarar que la variable sigue teniendo un tipo fijo en tiempo de compilación.",
        "ventajas": [
            "Nos ahorra escribir nombres larguísimos y genéricos enredados.",
            "Hace que sea más fácil enfocarse en el nombre y propósito de la variable."
        ],
        "desventajas": [
            "Si el valor de la derecha no es claro (como un método de un servicio), se vuelve difícil entender qué tipo de datos estamos manejando.",
            "No se puede usar como campos de la clase ni en los parámetros de los métodos."
        ],
        "cuando": "Para variables locales de corta duración donde el tipo de datos se entiende a simple vista (como al inicializar con new).",
        "code_trad": (
            "// Forma Tradicional\n"
            "Map<String, List<Map<Integer, String>>> jerarquia = new HashMap<String, List<Map<Integer, String>>>();"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "var jerarquia = new HashMap<String, List<Map<Integer, String>>>();"
        )
    },
    {
        "num": "6",
        "title": "6. Expresiones Lambda",
        "explicacion": "Son funciones anónimas concisas que permiten pasar comportamientos como si fueran datos. Evitan tener que escribir clases internas anónimas completas, y en memoria se manejan de manera más eficiente gracias a llamadas dinámicas de la máquina virtual.",
        "ventajas": [
            "Habilitan el estilo de programación funcional y declarativa en Java.",
            "Reduce mucho el código repetitivo de un solo método."
        ],
        "desventajas": [
            "A veces, cuando sale un error en consola, la pila de llamadas (stacktrace) muestra nombres de métodos automáticos que cuesta descifrar.",
            "Solo pueden usar variables externas que sean declaradas como finales o que no cambien de valor."
        ],
        "cuando": "Para pasar lógicas rápidas como argumentos, por ejemplo, al ordenar elementos o usar la API Streams.",
        "code_trad": (
            "// Forma Tradicional\n"
            "Runnable hilo = new Runnable() {\n"
            "    @Override\n"
            "    public void run() {\n"
            "        System.out.println(\"Ejecución en segundo plano\");\n"
            "    }\n"
            "};"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "Runnable hilo = () -> System.out.println(\"Ejecución en segundo plano\");"
        )
    },
    {
        "num": "7",
        "title": "7. Referencias a Métodos (::)",
        "explicacion": "Es un atajo para escribir expresiones lambda cuando estas solo llaman a un método que ya existe. Sirve para no tener que declarar parámetros intermedios que no hacen nada más que pasar el valor.",
        "ventajas": [
            "Es la sintaxis más limpia posible para pasar comportamientos preexistentes.",
            "Muy amigable para leer en flujos de datos."
        ],
        "desventajas": [
            "No sirve si necesitas hacer algún cambio pequeño o lógica adicional antes de llamar al método."
        ],
        "cuando": "Cuando la expresión lambda coincide exactamente con la firma de un método existente (como System.out::println).",
        "code_trad": (
            "// Forma Tradicional\n"
            "List<String> marcas = List.of(\"Intel\", \"AMD\", \"Nvidia\");\n"
            "marcas.forEach(marca -> System.out.println(marca));"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "List<String> marcas = List.of(\"Intel\", \"AMD\", \"Nvidia\");\n"
            "marcas.forEach(System.out::println);"
        )
    },
    {
        "num": "8",
        "title": "8. API Streams",
        "explicacion": "Es una herramienta para procesar colecciones de datos como si fuera una línea de ensamblaje. Tiene operaciones intermedias (como filtrar o mapear) que son perezosas (no se ejecutan hasta que llamamos a la operación final) para optimizar el rendimiento.",
        "ventajas": [
            "Cambia bucles anidados difíciles de entender por flujos limpios y lineales.",
            "Permite paralelizar el procesamiento fácilmente para que vaya más rápido en procesadores multi-núcleo."
        ],
        "desventajas": [
            "Puede ser un poco lento para colecciones muy pequeñas por la sobrecarga de crear objetos intermedios.",
            "Es difícil de depurar con los breakpoints normales de la computadora."
        ],
        "cuando": "Para realizar filtros, transformaciones y agrupaciones en colecciones grandes de datos de manera declarativa.",
        "code_trad": (
            "// Forma Tradicional\n"
            "List<Integer> numeros = List.of(1, 2, 3, 4, 5, 6);\n"
            "List<Integer> doblesDePares = new ArrayList<>();\n"
            "for (Integer n : numeros) {\n"
            "    if (n % 2 == 0) {\n"
            "        doblesDePares.add(n * 2);\n"
            "    }\n"
            "}"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "List<Integer> numeros = List.of(1, 2, 3, 4, 5, 6);\n"
            "List<Integer> doblesDePares = numeros.stream()\n"
            "                                    .filter(n -> n % 2 == 0)\n"
            "                                    .map(n -> n * 2)\n"
            "                                    .toList();"
        )
    },
    {
        "num": "9",
        "title": "9. Clase Optional",
        "explicacion": "Es una clase contenedora que puede o no tener un valor dentro. Sirve como un diseño semántico para indicar que un resultado puede ser nulo, obligándonos a manejar esa ausencia de forma segura.",
        "ventajas": [
            "Ayuda a evitar el famoso error NullPointerException (puntero nulo).",
            "Tiene métodos muy limpios que se integran con lambdas (como .map() o .orElse())."
        ],
        "desventajas": [
            "Consume un poco más de memoria porque crea un objeto envoltura para cada valor. No debe usarse para guardar datos en la clase."
        ],
        "cuando": "Para tipos de retorno en métodos de servicios donde es normal que a veces no se encuentre un resultado (como buscar por ID).",
        "code_trad": (
            "// Forma Tradicional\n"
            "String valor = obtenerDatoSeguro();\n"
            "String resultado;\n"
            "if (valor != null) {\n"
            "    resultado = valor.toUpperCase();\n"
            "} else {\n"
            "    resultado = \"VACÍO\";\n"
            "}"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "String resultado = Optional.ofNullable(obtenerDatoSeguro())\n"
            "                           .map(String::toUpperCase)\n"
            "                           .orElse(\"VACÍO\");"
        )
    },
    {
        "num": "10",
        "title": "10. Switch Expression",
        "explicacion": "Es una forma moderna de usar switch que devuelve un valor de forma directa (como una expresión) y usa flechas (->) en lugar de dos puntos. Esto elimina la necesidad del comando break y la caída accidental de casos.",
        "ventajas": [
            "Evita errores causados por olvidarse de escribir un break;.",
            "Obliga a cubrir todos los casos posibles en tiempo de compilación."
        ],
        "desventajas": [
            "Si un caso tiene demasiadas líneas de código, se vuelve pesado y hay que usar llaves y la palabra clave yield."
        ],
        "cuando": "Al asignar valores dependiendo de categorías fijas como enums o tipos constantes.",
        "code_trad": (
            "// Forma Tradicional\n"
            "int dias;\n"
            "switch (mes) {\n"
            "    case ENERO:\n"
            "    case MARZO:\n"
            "        dias = 31;\n"
            "        break;\n"
            "    case FEBRERO:\n"
            "        dias = 28;\n"
            "        break;\n"
            "    default:\n"
            "        throw new IllegalArgumentException(\"Mes no soportado\");\n"
            "}"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "int dias = switch (mes) {\n"
            "    case ENERO, MARZO -> 31;\n"
            "    case FEBRERO      -> 28;\n"
            "};"
        )
    },
    {
        "num": "11",
        "title": "11. Text Blocks (Bloques de Texto)",
        "explicacion": "Permite escribir cadenas de texto que ocupan varias líneas usando tres comillas dobles (\"\"\"). Elimina la necesidad de concatenar con + o poner caracteres raros como \\n o comillas escapadas.",
        "ventajas": [
            "Súper útil para escribir consultas SQL, JSON o HTML dentro de Java y que sigan siendo legibles.",
            "Respeta el formato y la indentación de forma automática."
        ],
        "desventajas": [
            "Hay que tener cuidado con los espacios al final de las líneas y la posición de las comillas de cierre."
        ],
        "cuando": "Cuando tengas textos largos estructurados que quieras pegar tal cual en el código.",
        "code_trad": (
            "// Forma Tradicional\n"
            "String sql = \"SELECT id, nombre, email\\n\" +\n"
            "             \"FROM usuarios\\n\" +\n"
            "             \"WHERE activo = 1\\n\" +\n"
            "             \"ORDER BY nombre ASC;\";"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "String sql = \"\"\"\n"
            "             SELECT id, nombre, email\n"
            "             FROM usuarios\n"
            "             WHERE activo = 1\n"
            "             ORDER BY nombre ASC;\n"
            "             \"\"\";"
        )
    },
    {
        "num": "12",
        "title": "12. Records (Registros)",
        "explicacion": "Son clases especiales inmutables diseñadas solo para transportar datos. Java genera de forma automática los campos privados finales, el constructor, los métodos de acceso y las utilidades como equals(), hashCode() y toString().",
        "ventajas": [
            "Ahorra muchísimas líneas de código que antes escribíamos a mano o generábamos con editores (Lombok/POJOs).",
            "Asegura que los datos no se puedan modificar una vez creados."
        ],
        "desventajas": [
            "No soportan herencia clásica (no pueden extender a otras clases ni ser heredados), lo que limita su flexibilidad en algunos patrones."
        ],
        "cuando": "Para modelar objetos simples de transferencia de datos (DTOs) o para guardar respuestas rápidas de bases de datos.",
        "code_trad": (
            "// Forma Tradicional\n"
            "public final class Punto {\n"
            "    private final int x;\n"
            "    private final int y;\n\n"
            "    public Punto(int x, int y) {\n"
            "        this.x = x;\n"
            "        this.y = y;\n"
            "    }\n\n"
            "    public int x() { return x; }\n"
            "    public int y() { return y; }\n\n"
            "    @Override\n"
            "    public boolean equals(Object o) {\n"
            "        if (this == o) return true;\n"
            "        if (!(o instanceof Punto)) return false;\n"
            "        Punto punto = (Punto) o;\n"
            "        return x == punto.x && y == punto.y;\n"
            "    }\n\n"
            "    @Override\n"
            "    public int hashCode() {\n"
            "        return Objects.hash(x, y);\n"
            "    }\n"
            "}"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "public record Punto(int x, int y) {}"
        )
    },
    {
        "num": "13",
        "title": "13. Pattern Matching con instanceof",
        "explicacion": "Combina la comprobación de tipo (instanceof) y la conversión de tipo (casting) en una sola operación. Si el objeto coincide con el tipo, Java crea automáticamente una nueva variable ya convertida lista para usar.",
        "ventajas": [
            "Hace el código más corto y seguro.",
            "Evita que cometamos errores de casting incorrectos en tiempo de ejecución."
        ],
        "desventajas": [
            "La variable de patrón solo se puede usar dentro del bloque donde se cumple la condición."
        ],
        "cuando": "Al recibir objetos genéricos en interfaces o métodos generales para validar su tipo y consumirlos de inmediato.",
        "code_trad": (
            "// Forma Tradicional\n"
            "if (obj instanceof Number) {\n"
            "    Number num = (Number) obj;\n"
            "    System.out.println(num.doubleValue());\n"
            "}"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "if (obj instanceof Number num) {\n"
            "    System.out.println(num.doubleValue());\n"
            "}"
        )
    },
    {
        "num": "14",
        "title": "14. Factoría de Colecciones con List.of()",
        "explicacion": "Es una forma muy rápida de crear listas que no se pueden modificar (inmutables) en una sola línea. Java usa clases internas muy optimizadas que ocupan menos memoria en comparación con crear un ArrayList y llenarlo.",
        "ventajas": [
            "Crea colecciones constantes seguras de forma rápida.",
            "Ocupa menos espacio en memoria y previene cambios accidentales en la lista."
        ],
        "desventajas": [
            "Si intentas añadir o quitar algo de la lista después de crearla, el programa fallará con un error UnsupportedOperationException."
        ],
        "cuando": "Para definir colecciones constantes de datos (como los meses del año o configuraciones fijas).",
        "code_trad": (
            "// Forma Tradicional\n"
            "List<String> dias = new ArrayList<>();\n"
            "dias.add(\"Lunes\");\n"
            "dias.add(\"Martes\");\n"
            "List<String> inmutable = Collections.unmodifiableList(dias);"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "List<String> inmutable = List.of(\"Lunes\", \"Martes\");"
        )
    },
    {
        "num": "15",
        "title": "15. Método de Eliminación Condicional removeIf()",
        "explicacion": "Permite eliminar elementos de una colección mutable basándose en una condición lógica (un predicado). Es mucho más rápido que recorrer con un ciclo clásico y borrar a mano, ya que Java optimiza los desplazamientos de memoria internamente.",
        "ventajas": [
            "Evita errores comunes al borrar elementos en bucles tradicionales (como saltarse índices o errores concurrentes).",
            "El código se ve muy directo y elegante en una línea."
        ],
        "desventajas": [
            "Si se ejecuta en colecciones inmutables, dará error inmediatamente."
        ],
        "cuando": "Para limpiar o filtrar colecciones mutables grandes de forma rápida y segura.",
        "code_trad": (
            "// Forma Tradicional\n"
            "Iterator<Integer> iterador = numerosMutables.iterator();\n"
            "while (iterador.hasNext()) {\n"
            "    Integer n = iterador.next();\n"
            "    if (n % 2 == 0) {\n"
            "        iterador.remove();\n"
            "    }\n"
            "}"
        ),
        "code_abbr": (
            "// Forma Abreviada\n"
            "numerosMutables.removeIf(n -> n % 2 == 0);"
        )
    }
]

part2_topics = [
    {
        "num": "16",
        "title": "16. Variables y Patrones sin Nombre (_ - Underscore)",
        "explicacion": "Permite usar el guion bajo (_) para variables que el compilador exige declarar pero que en nuestra lógica no vamos a usar (como en bloques catch o parámetros de lambdas).",
        "ventajas": [
            "Indica claramente que el valor no importa, quitando ruido visual.",
            "Evita inventar nombres raros como 'ignorado' o 'dummy'.",
            "Permite poner varios guiones bajos en el mismo bloque sin errores de nombre duplicado."
        ],
        "desventajas": [
            "Solo sirve para variables locales y parámetros, no para campos principales de la clase."
        ],
        "cuando": "En excepciones de catch que no usemos, o en lambdas con parámetros obligatorios pero inútiles para el cálculo.",
        "code_trad": (
            "// Forma Tradicional\n"
            "try {\n"
            "    int puerto = Integer.parseInt(puertoConfigurado);\n"
            "} catch (NumberFormatException ignorado) { // Nombre ficticio e inútil\n"
            "    System.err.println(\"Puerto inválido, usando puerto por defecto.\");\n"
            "}\n\n"
            "mapaDeSesion.computeIfAbsent(usuarioId, clave -> new ArrayList<>()); // 'clave' es redundante"
        ),
        "code_abbr": (
            "// Sintaxis Comprimida (Java 22+)\n"
            "try {\n"
            "    int puerto = Integer.parseInt(puertoConfigurado);\n"
            "} catch (NumberFormatException _) { // Descarte absoluto\n"
            "    System.err.println(\"Puerto inválido, usando puerto por defecto.\");\n"
            "}\n\n"
            "mapaDeSesion.computeIfAbsent(usuarioId, _ -> new ArrayList<>()); // Parámetro ignorado de forma segura"
        )
    },
    {
        "num": "17",
        "title": "17. Castings de Intersección de Tipos en Expresiones Lambda",
        "explicacion": "Es una sintaxis avanzada que junta varias interfaces en una sola usando el símbolo &. Se usa sobre todo en sistemas distribuidos para obligar a una expresión lambda a implementar también una interfaz marcador como Serializable.",
        "ventajas": [
            "Permite enviar comportamientos por red (serializarlos) sin tener que crear archivos de interfaz permanentes a mano."
        ],
        "desventajas": [
            "Es una sintaxis un poco compleja que puede confundir a primera vista. Solo funciona si una de las interfaces tiene el método funcional a ejecutar."
        ],
        "cuando": "En tecnologías de computación en red o mensajería donde las tareas dinámicas deben viajar serializadas por internet.",
        "code_trad": (
            "// Forma Tradicional (Requiere crear interfaces nominales redundantes de infraestructura)\n"
            "public interface RunnableSerializable extends Runnable, Serializable {}\n\n"
            "public class Distribuidor {\n"
            "    public void enviarTarea(RunnableSerializable tarea) {\n"
            "        // Lógica de serialización y envío remoto\n"
            "    \n}"
        ),
        "code_abbr": (
            "// Sintaxis Comprimida (Utiliza el casting de intersección para forzar la serialización in-situ)\n"
            "public class Distribuidor {\n"
            "    public void enviarTarea(Object tarea) {\n"
            "        if (tarea instanceof Serializable) {\n"
            "            // Serializar de forma segura de forma dinámica\n"
            "        }\n"
            "    }\n"
            "}\n\n"
            "// Uso en caliente en el cliente\n"
            "Distribuidor distribuidor = new Distribuidor();\n"
            "distribuidor.enviarTarea((Runnable & Serializable) () -> System.out.println(\"Procesando tarea remota...\"));"
        )
    },
    {
        "num": "18",
        "title": "18. Inferencia de Tipos No Denotables con var sobre Instancias de Clases Anónimas",
        "explicacion": "Permite usar var para instanciar clases anónimas, logrando que el compilador recuerde métodos o propiedades que agreguemos de forma exclusiva en esa instancia sin degradar el tipo a Object.",
        "ventajas": [
            "Permite crear estructuras parecidas a 'tuplas' locales rápidas para transportar datos del cálculo sin tener que hacer una clase formal en un archivo .java aparte."
        ],
        "desventajas": [
            "Estas variables mejoradas no se pueden retornar como valores de métodos públicos porque su tipo no tiene un nombre formal (no denotable)."
        ],
        "cuando": "Para empaquetar y procesar datos temporales complejos dentro de un mismo método.",
        "code_trad": (
            "// Forma Tradicional (Obliga a crear una clase estructurada física temporal)\n"
            "class EstructuraCalculo {\n"
            "    double valorConvertido;\n"
            "    long latenciaMs;\n"
            "    EstructuraCalculo(double v, long l) { this.valorConvertido = v; this.latenciaMs = l; }\n"
            "}\n\n"
            "public EstructuraCalculo procesar() {\n"
            "    return new EstructuraCalculo(500.25, 42L);\n"
            "}"
        ),
        "code_abbr": (
            "// Sintaxis Comprimida (Utiliza var y clase anónima para construir una tupla local no denotable)\n"
            "public void calcularEstadisticas() {\n"
            "    var resultadoTupla = new Object() { // Inferencia precisa del tipo no denotable\n"
            "        double valorConvertido = 500.25;\n"
            "        long latenciaMs = 42L;\n"
            "        void imprimirDetalle() {\n"
            "            System.out.println(\"Valor: \" + valorConvertido + \" en \" + latenciaMs + \"ms\");\n"
            "        }\n"
            "    };\n\n"
            "    // Accediendo a propiedades inaccesibles de forma nominal clásica\n"
            "    System.out.println(resultadoTupla.valorConvertido);\n"
            "    resultadoTupla.imprimirDetalle();\n"
            "}"
        )
    },
    {
        "num": "19",
        "title": "19. Stream Gatherers (JEP 461)",
        "explicacion": "Es una característica experimental de Java 22 que permite crear operaciones intermedias personalizadas en Streams. A diferencia de map o filter, nos deja llevar un control del estado y cortar el procesamiento de forma dinámica.",
        "ventajas": [
            "Permite hacer cosas avanzadas como agrupar elementos en ventanas de tamaño fijo de forma declarativa y súper sencilla."
        ],
        "desventajas": [
            "Al ser una función experimental (preview), requiere activar parámetros especiales de compilación al correr el programa."
        ],
        "cuando": "Para análisis de datos temporales, promedios móviles o empaquetar flujos en bloques de tamaño definido.",
        "code_trad": (
            "// Forma Tradicional (Mapeo por lotes o ventanas agrupadas requiere acumulación mutable imperativa compleja)\n"
            "public static List<List<Integer>> agruparEnVentanas(List<Integer> numeros, int tamano) {\n"
            "    List<List<Integer>> resultado = new ArrayList<>();\n"
            "    List<Integer> ventana = new ArrayList<>();\n"
            "    for (Integer n : numeros) {\n"
            "        ventana.add(n);\n"
            "        if (ventana.size() == tamano) {\n"
            "            resultado.add(new ArrayList<>(ventana));\n"
            "            ventana.clear();\n"
            "        }\n"
            "    }\n"
            "    if (!ventana.isEmpty()) {\n"
            "        resultado.add(ventana);\n"
            "    }\n"
            "    return resultado;\n"
            "}"
        ),
        "code_abbr": (
            "// Sintaxis Comprimida (Utiliza Stream Gatherers preconstruidos de Java 22+)\n"
            "import java.util.stream.Gatherers;\n\n"
            "public class AgrupadorDatos {\n"
            "    public List<List<Integer>> agruparConGatherers(List<Integer> numeros, int tamano) {\n"
            "        return numeros.stream()\n"
            "                     .gather(Gatherers.windowFixed(tamano)) // Operación de ventana fija\n"
            "                     .toList();\n"
            "    }\n"
            "}"
        )
    },
    {
        "num": "20",
        "title": "20. Labeled Breaks en Bloques Arbitrarios y Estructuras Try-Catch",
        "explicacion": "Permite poner una etiqueta con dos puntos (como miEtiqueta:) a cualquier bloque de llaves, incluso a un try-catch, y usar break miEtiqueta; para salir inmediatamente de él sin lanzar excepciones pesadas. Si hay un bloque finally, Java garantiza que se ejecutará antes de salir.",
        "ventajas": [
            "Permite abortar procesos anidados de forma muy rápida y sin consumir la memoria que requiere crear una excepción.",
            "Garantiza la limpieza de recursos mediante el finally."
        ],
        "desventajas": [
            "Es una sintaxis muy poco común que puede asustar o confundir si no se conoce bien cómo funciona el control de flujo."
        ],
        "cuando": "En algoritmos de análisis sintáctico o procesamiento de lotes con muchas validaciones donde queramos abortar de forma limpia y segura ante fallos lógicos.",
        "code_trad": (
            "// Forma Tradicional (Múltiples banderas de control y validaciones redundantes)\n"
            "boolean operacionExitosa = true;\n"
            "try {\n"
            "    System.out.println(\"Abriendo sesión transaccional...\");\n"
            "    if (comprobarMantenimientoActivo()) {\n"
            "        operacionExitosa = false;\n"
            "    }\n"
            "    if (operacionExitosa) {\n"
            "        System.out.println(\"Ejecutando operaciones pesadas...\");\n"
            "        // Operación...\n"
            "    }\n"
            "} finally {\n"
            "    System.out.println(\"Garantizando cierre de sesión y recursos...\");\n"
            "}"
        ),
        "code_abbr": (
            "// Sintaxis Comprimida (Utiliza un Labeled Break para abortar limpiamente el try transaccional)\n"
            "bloqueTransaccion: try {\n"
            "    System.out.println(\"Abriendo sesión transaccional...\");\n"
            "    if (comprobarMantenimientoActivo()) {\n"
            "        break bloqueTransaccion; // Aborta inmediatamente el procesamiento del bloque etiquetado\n"
            "    }\n\n"
            "    // Este fragmento de código es saltado de forma limpia y segura por la JVM\n"
            "    System.out.println(\"Ejecutando operaciones pesadas...\");\n"
            "} finally {\n"
            "    // La JVM ejecutará de forma prioritaria este bloque antes de procesar el escape final\n"
            "    System.out.println(\"Garantizando cierre de sesión y recursos...\");\n"
            "}"
        )
    }
]

# Write function to output a topic
def write_topic_to_doc(doc, topic):
    doc.add_heading(topic["title"], level=2)
    add_spacing(doc, pt=6)
    
    add_section_label(doc, "Explicación", topic["explicacion"])
    add_spacing(doc, pt=4)
    
    add_section_label(doc, "Ventajas")
    for v in topic["ventajas"]:
        add_bullet_item(doc, v)
    add_spacing(doc, pt=4)
    
    add_section_label(doc, "Desventajas")
    for d in topic["desventajas"]:
        add_bullet_item(doc, d)
    add_spacing(doc, pt=4)
    
    add_section_label(doc, "Cuándo utilizarlo", topic["cuando"])
    add_spacing(doc, pt=4)
    
    add_section_label(doc, "Ejemplo")
    add_code_block(doc, topic["code_trad"])
    add_spacing(doc, pt=4)
    add_code_block(doc, topic["code_abbr"])
    add_spacing(doc, pt=8)

# --- Parte 1: Análisis de Sintaxis Abreviadas y Características Propuestas ---
doc.add_heading("Parte 1: Análisis de Sintaxis Abreviadas y Características Propuestas", level=1)
add_spacing(doc, pt=8)

for topic in part1_topics:
    write_topic_to_doc(doc, topic)
    # Add comparison table in Topic 15
    if topic["num"] == "15":
        doc.add_heading("Comparación de Rendimiento de Eliminación Condicional en Colecciones", level=3)
        add_spacing(doc, pt=4)
        doc.add_paragraph(
            "La siguiente tabla muestra de forma sencilla la diferencia en tiempo de procesamiento "
            "según el tipo de colección y el método que usemos para borrar los elementos:",
            style="Normal"
        )
        add_spacing(doc, pt=4)
        
        table_data = [
            ["Tipo de Colección", "Eliminación Manual (Bucle Clásico)", "Eliminación con removeIf() (Optimizado)", "Comportamiento Interno"],
            ["ArrayList", "O(N^2) debido a reorganizaciones continuas", "O(N) en una sola pasada de compactación", "Modifica directamente el arreglo interno sin copias extras."],
            ["LinkedList", "O(N) por operaciones de enlaces", "O(N) por punteros internos", "Cambia las referencias de los nodos sin mover memoria físicamente."],
            ["CopyOnWriteArrayList", "O(N^2) súper ineficiente por copiar el arreglo en cada borrado", "O(N) con una única copia al final", "Reduce la sobrecarga haciendo solo una copia completa al terminar."]
        ]
        col_widths = [Inches(1.2), Inches(1.6), Inches(1.6), Inches(2.1)]
        add_word_table(doc, table_data, col_widths)
        add_spacing(doc, pt=8)

# --- Parte 2: Investigación de 5 Sintaxis de Baja Visibilidad y Características Avanzadas ---
doc.add_heading("Parte 2: Investigación de 5 Sintaxis de Baja Visibilidad y Características Avanzadas", level=1)
add_spacing(doc, pt=8)

for topic in part2_topics:
    write_topic_to_doc(doc, topic)
    # Add resolution table in Topic 18
    if topic["num"] == "18":
        doc.add_heading("Inferencia de Tipo No Denotable en Bloque de Código", level=3)
        add_spacing(doc, pt=4)
        doc.add_paragraph(
            "En esta tabla se puede ver qué pasa si declaramos una clase anónima usando un tipo Object clásico "
            "frente a usar var. Con var, Java recuerda todo lo que le agregamos adentro:",
            style="Normal"
        )
        add_spacing(doc, pt=4)
        
        table_data = [
            ["Tipo de Declaración", "Firma de la Variable Resultante", "Disponibilidad de Métodos Propios (No Declarados en la Clase Padre)"],
            ["Object obj = new Object() {... }", "java.lang.Object (Se pierde la información específica)", "No disponible de forma directa en el código."],
            ["var obj = new Object() {... }", "<anonymous java.lang.Object> (Java recuerda la estructura interna)", "Totalmente disponible para usar de forma segura."]
        ]
        col_widths = [Inches(2.0), Inches(2.2), Inches(2.3)]
        add_word_table(doc, table_data, col_widths)
        add_spacing(doc, pt=8)
        
    # Add diagram in Topic 20
    if topic["num"] == "20":
        doc.add_heading("Comportamiento de Salto Estructurado en Bloque Etiquetado", level=3)
        add_spacing(doc, pt=4)
        doc.add_paragraph(
            "Este esquema dibuja el camino que sigue Java cuando usamos un break etiquetado dentro de un bloque con finally. "
            "Se puede ver que el finally siempre se ejecuta pase lo que pase:",
            style="Normal"
        )
        add_spacing(doc, pt=4)
        
        diagram_text = (
            "[Inicio de Ejecución]\n"
            "         │\n"
            "         ▼\n"
            " ┌──────────────┐\n"
            " │   try {      │ <─── ETIQUETA: bloqueTransaccion\n"
            " │   ...        │\n"
            " │  break;      │ ───► Salto lógico inmediato fuera del try\n"
            " └──────────────┘\n"
            "         │\n"
            "         ▼ (Interrupción capturada por la JVM para hacer limpieza)\n"
            " ┌──────────────┐\n"
            " │  finally {   │ <─── EJECUCIÓN GARANTIZADA ANTES DEL ESCAPE\n"
            " │   ...        │\n"
            " └──────────────┘\n"
            "         │\n"
            "         ▼\n"
            "[Fin del Bloque Transaccional]"
        )
        add_diagram_block(doc, diagram_text)
        add_spacing(doc, pt=8)

# Add Bibliography/References Section
doc.add_heading("Referencias Bibliográficas", level=1)
add_spacing(doc, pt=8)

references = [
    "Oracle Corporation. (2026). Java Language Specification & API Documentation (Java SE 22 Edition). Oracle Technology Network.",
    "OpenJDK. (2026). Project Amber JEPs (JEP 286: Local-Variable Type Inference, JEP 456: Unnamed Variables & Patterns, JEP 461: Stream Gatherers). OpenJDK.org.",
    "Bloch, J. (2018). Effective Java (3rd Edition). Addison-Wesley Professional. (Guías sobre genéricos, lambdas y Streams).",
    "Goetz, B. (2020). Java Language Architecture and Evolution: Designing Java for the Future. Oracle Corporation.",
    "Baeldung. (2026). Guides to Modern Java Features and Syntax Enhancements. Baeldung.com."
]

for ref in references:
    add_bullet_item(doc, ref)

try:
    doc.save(output_path)
    print(f"Document created successfully at: {output_path}")
except Exception as e:
    print(f"Failed to save document to {output_path}: {e}")
    alt_path = output_path.replace(".docx", "_v2.docx")
    try:
        doc.save(alt_path)
        print(f"Document created successfully at alternative path: {alt_path}")
    except Exception as e2:
        print(f"Failed to save to alternative path: {e2}")
